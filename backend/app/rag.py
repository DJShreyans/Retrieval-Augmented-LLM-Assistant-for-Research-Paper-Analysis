import os
import re
import requests
import fitz  # PyMuPDF
import docx  # python-docx
import legacy_doc  # legacy-doc
from bs4 import BeautifulSoup  # beautifulsoup4
import chromadb
from chromadb.utils import embedding_functions

# Setup persistent directory paths relative to this file
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DB_PATH = os.path.join(BASE_DIR, "data", "chromadb")
UPLOADS_PATH = os.path.join(BASE_DIR, "data", "uploads")

# Ensure directories exist
os.makedirs(DB_PATH, exist_ok=True)
os.makedirs(UPLOADS_PATH, exist_ok=True)

from chromadb.utils.embedding_functions import DefaultEmbeddingFunction

# Initialize ChromaDB client and embedding model (ONNX-based all-MiniLM-L6-v2)
client = chromadb.PersistentClient(path=DB_PATH)
embedding_func = DefaultEmbeddingFunction()
collection = client.get_or_create_collection(
    name="researchmate_docs", 
    embedding_function=embedding_func
)



def extract_text_from_pdf(file_path: str) -> str:
    """Extracts text content from a PDF document using PyMuPDF (fitz)."""
    text = ""
    try:
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text() + "\n"
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
        raise e
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extracts text content from a DOCX document using python-docx."""
    text = ""
    try:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                text += " | ".join(row_text) + "\n"
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
        raise e
    return text

def extract_text_from_doc(file_path: str) -> str:
    """Extracts text content from a legacy Word .doc file using legacy-doc."""
    try:
        with open(file_path, "rb") as f:
            result = legacy_doc.extract_text(f.read())
            return result.text
    except Exception as e:
        print(f"Error reading legacy DOC {file_path}: {e}")
        raise e

def ingest_url(url: str) -> dict:
    """
    Scrapes website text content, clean scripts/boilerplates, 
    splits it into chunks, and vector-indexes it in ChromaDB.
    """
    try:
        # Validate URL starts with http/https
        if not url.startswith(("http://", "https://")):
            raise ValueError("URL must start with 'http://' or 'https://'")
            
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        }
        
        response = requests.get(url, headers=headers, timeout=12)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, "html.parser")
        
        # Strip out code elements, styles, navs, footers, scripts
        for element in soup(["script", "style", "nav", "footer", "header", "aside", "form"]):
            element.decompose()
            
        # Get page title or fall back to URL
        title = soup.title.string.strip() if (soup.title and soup.title.string) else url
        
        # Extract text body and clean spacing
        raw_text = soup.get_text(separator=" ")
        text = re.sub(r'\s+', ' ', raw_text).strip()
        
        chunks = split_text_into_chunks(text)
        if not chunks:
            chunks = [f"No readable content found on website page: {url}"]
            
        ids = [f"{url}_chunk_{i}" for i in range(len(chunks))]
        metadatas = [{"source": url, "chunk_index": i} for i in range(len(chunks))]
        
        # Store chunks in vector database
        collection.add(
            ids=ids,
            documents=chunks,
            metadatas=metadatas
        )
        
        return {
            "filename": url,
            "title": title,
            "num_chunks": len(chunks),
            "total_characters": len(text)
        }
    except Exception as e:
        print(f"Error scraping website page {url}: {e}")
        raise e

def split_text_into_chunks(text: str, chunk_size_words: int = 500, overlap_words: int = 75) -> list[str]:
    """
    Splits text into chunks of roughly 500-800 tokens.
    Uses words as an approximation (1 token ≈ 0.75 words, so 500-800 tokens is roughly 375-600 words).
    We chunk by sentences and pack them to maintain logical boundaries.
    """
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    
    # Split text into sentences using simple regex
    sentences = re.split(r'(?<=[.!?])\s+', text)
    
    chunks = []
    current_chunk = []
    current_word_count = 0
    
    for sentence in sentences:
        sentence_words = sentence.split()
        sentence_word_count = len(sentence_words)
        
        if not sentence_words:
            continue
            
        # If adding this sentence exceeds our threshold and we already have some text
        if current_word_count + sentence_word_count > chunk_size_words and current_chunk:
            # Join and add current chunk
            chunks.append(" ".join(current_chunk))
            
            # Start a new chunk, carrying over some words for overlap
            # Find the index to slice from based on overlap word count
            overlap_sentences = []
            overlap_count = 0
            for s in reversed(current_chunk):
                s_words = s.split()
                if overlap_count + len(s_words) <= overlap_words:
                    overlap_sentences.insert(0, s)
                    overlap_count += len(s_words)
                else:
                    break
            
            current_chunk = overlap_sentences + [sentence]
            current_word_count = overlap_count + sentence_word_count
        else:
            current_chunk.append(sentence)
            current_word_count += sentence_word_count
            
    if current_chunk:
        chunks.append(" ".join(current_chunk))
        
    return chunks

def ingest_document(file_name: str, file_path: str) -> dict:
    """
    Extracts text, splits it into chunks, generates embeddings,
    and indexes them into ChromaDB.
    """
    ext = os.path.splitext(file_name)[1].lower()
    
    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif ext == ".docx":
        text = extract_text_from_docx(file_path)
    elif ext == ".doc":
        text = extract_text_from_doc(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
        
    chunks = split_text_into_chunks(text)
    
    if not chunks:
        # Fallback if the document is empty or text extraction failed
        chunks = [f"This document is empty. Filename: {file_name}"]
        
    ids = [f"{file_name}_chunk_{i}" for i in range(len(chunks))]
    metadatas = [{"source": file_name, "chunk_index": i} for i in range(len(chunks))]
    
    # Store chunks into ChromaDB
    # Collection will handle embedding generation via embedding_func automatically
    collection.add(
        ids=ids,
        documents=chunks,
        metadatas=metadatas
    )
    
    return {
        "filename": file_name,
        "num_chunks": len(chunks),
        "total_characters": len(text)
    }

# Reranking is performed via NVIDIA cloud reranking API to avoid loading local models
def query_vector_store(query_text: str, file_filter: list[str] = None, top_k: int = 4, rerank: bool = True) -> list[dict]:
    """
    Queries ChromaDB for the most relevant document chunks matching a query.
    Can filter results to specific uploaded document names.
    If rerank is True, retrieves candidates and re-ranks them using NVIDIA cloud reranking.
    """
    where_clause = {}
    if file_filter:
        if len(file_filter) == 1:
            where_clause = {"source": file_filter[0]}
        else:
            where_clause = {"$or": [{"source": f} for f in file_filter]}
            
    # Retrieve more candidate chunks if we are re-ranking
    query_k = top_k * 3 if rerank else top_k
    
    results = collection.query(
        query_texts=[query_text],
        n_results=query_k,
        where=where_clause if file_filter else None
    )
    
    formatted_results = []
    if results and results["documents"]:
        docs = results["documents"][0]
        metas = results["metadatas"][0]
        ids = results["ids"][0]
        distances = results["distances"][0] if "distances" in results else [0]*len(docs)
        
        for i in range(len(docs)):
            formatted_results.append({
                "id": ids[i],
                "content": docs[i],
                "source": metas[i]["source"],
                "chunk_index": metas[i]["chunk_index"],
                "score": 1 - distances[i]  # Convert distance to similarity score
            })
            
    # Re-rank candidate results using Cloud-hosted NVIDIA Reranker
    if rerank and len(formatted_results) > 1:
        nvidia_key = os.environ.get("NVIDIA_API_KEY", "").strip()
        if nvidia_key:
            try:
                headers = {
                    "Authorization": f"Bearer {nvidia_key}",
                    "Content-Type": "application/json"
                }
                payload = {
                    "model": "nvidia/rerank-qa-mistral-4b",
                    "query": {
                        "text": query_text
                    },
                    "passages": [
                        {"text": item["content"]} for item in formatted_results
                    ]
                }
                response = requests.post(
                    "https://ai.api.nvidia.com/v1/retrieval/nvidia/reranking",
                    headers=headers,
                    json=payload,
                    timeout=12
                )
                # Fallback to model version suffix if needed
                if response.status_code != 200:
                    payload["model"] = "nv-rerank-qa-mistral-4b:1"
                    response = requests.post(
                        "https://ai.api.nvidia.com/v1/retrieval/nvidia/reranking",
                        headers=headers,
                        json=payload,
                        timeout=12
                    )
                
                if response.status_code == 200:
                    res_data = response.json()
                    for rank_item in res_data.get("data", []):
                        idx = rank_item["index"]
                        score = rank_item.get("logit", rank_item.get("score", 0.0))
                        formatted_results[idx]["rerank_score"] = float(score)
                        formatted_results[idx]["score"] = float(score)
                    
                    # Sort by re-ranked scores (descending)
                    formatted_results.sort(key=lambda x: x.get("rerank_score", 0.0), reverse=True)
                else:
                    print(f"[Reranker] Cloud API error: {response.text}")
            except Exception as e:
                print(f"[Reranker] Failed to run cloud-hosted re-ranking: {e}")
            
    return formatted_results[:top_k]


def delete_document_from_vector_store(file_name: str):
    """Deletes all vectorized chunks associated with a document from ChromaDB."""
    collection.delete(where={"source": file_name})
